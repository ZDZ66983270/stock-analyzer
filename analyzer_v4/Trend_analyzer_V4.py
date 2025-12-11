import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
import re


def detect_macd_divergence(df: pd.DataFrame, period: int = 60):
    signals = []
    df = df.tail(period).copy()
    df['MACD'] = df['MACD'].astype(float)
    df['收盘'] = df['收盘'].astype(float)

    for i in range(2, len(df) - 1):
        prev_close_peak = df['收盘'].iloc[i - 2]
        curr_close_peak = df['收盘'].iloc[i]
        prev_macd_peak = df['MACD'].iloc[i - 2]
        curr_macd_peak = df['MACD'].iloc[i]

        if curr_close_peak > prev_close_peak and curr_macd_peak < prev_macd_peak:
            signals.append((df['时间'].iloc[i], 'MACD顶背离'))
        elif curr_close_peak < prev_close_peak and curr_macd_peak > prev_macd_peak:
            signals.append((df['时间'].iloc[i], 'MACD底背离'))

    return signals


def detect_kdj_divergence(df: pd.DataFrame, period: int = 60):
    signals = []
    df = df.tail(period).copy()
    df['J'] = df['J'].astype(float)
    df['收盘'] = df['收盘'].astype(float)

    for i in range(2, len(df) - 1):
        prev_close_peak = df['收盘'].iloc[i - 2]
        curr_close_peak = df['收盘'].iloc[i]
        prev_j_peak = df['J'].iloc[i - 2]
        curr_j_peak = df['J'].iloc[i]

        if curr_close_peak > prev_close_peak and curr_j_peak < prev_j_peak:
            signals.append((df['时间'].iloc[i], 'KDJ顶背离'))
        elif curr_close_peak < prev_close_peak and curr_j_peak > prev_j_peak:
            signals.append((df['时间'].iloc[i], 'KDJ底背离'))

    return signals


def detect_rsi_divergence(df, period=14, lookback=60):
    """
    检测RSI顶/底背离，返回信号列表（时间戳, 类型）
    """
    signals = []
    if '收盘' not in df.columns or f'RSI{period}' not in df.columns:
        return signals

    close = df['收盘']
    rsi = df[f'RSI{period}']

    for i in range(lookback, len(df)):
        # 顶背离：价格创新高，RSI未创新高
        price_window = close[i-lookback:i+1]
        rsi_window = rsi[i-lookback:i+1]
        if close[i] == price_window.max() and rsi[i] < rsi_window.max():
            signals.append((df['时间'].iloc[i], 'RSI顶背离'))
        # 底背离：价格创新低，RSI未创新低
        if close[i] == price_window.min() and rsi[i] > rsi_window.min():
            signals.append((df['时间'].iloc[i], 'RSI底背离'))
    return signals


def detect_ma_divergence(df, ma_col='MA20', lookback=60):
    """
    检测均线顶/底背离，返回信号列表（时间戳, 类型）
    """
    signals = []
    if '收盘' not in df.columns or ma_col not in df.columns:
        return signals

    close = df['收盘']
    ma = df[ma_col]

    for i in range(lookback, len(df)):
        price_window = close[i-lookback:i+1]
        ma_window = ma[i-lookback:i+1]
        if close[i] == price_window.max() and ma[i] < ma_window.max():
            signals.append((df['时间'].iloc[i], f'{ma_col}顶背离'))
        if close[i] == price_window.min() and ma[i] > ma_window.min():
            signals.append((df['时间'].iloc[i], f'{ma_col}底背离'))
    return signals


def filter_recent_signals(signals, df, days=7, bars=20):
    if not signals:
        return []
    # 取最新时间
    last_time = df['时间'].iloc[-1]
    # 只保留最近7天
    signals = [s for s in signals if pd.to_datetime(s[0]) >= pd.to_datetime(last_time) - pd.Timedelta(days=days)]
    # 只保留最近20根K线
    last_indices = set(df.index[-bars:])
    signals = [s for s in signals if df[df['时间'] == s[0]].index[0] in last_indices]
    return signals


def add_signals_to_docx(doc, signals, title):
    doc.add_paragraph(f"【{title}】")
    if not signals:
        doc.add_paragraph(f"{title} 无信号")
        return

    for signal in signals:
        p = doc.add_paragraph()
        if len(signal) >= 4:  # 新格式：时间、类型、强度、描述、详情
            time, type_, strength, description, details = signal
            run = p.add_run(f"{time} - {type_} ({strength})")
            if strength == "强":
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
            elif strength == "弱":
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # 绿色
            p.add_run(f"\n{description}")
        else:  # 旧格式：时间、类型
            time, type_ = signal
            run = p.add_run(f"{time} - {type_}")
            if '顶背离' in type_:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
            elif '底背离' in type_:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色


def analyze_fund_flow(fund_flow_path):
    import pandas as pd
    from datetime import datetime

    if not os.path.exists(fund_flow_path):
        return "未找到资金流向数据文件。"

    df = pd.read_excel(fund_flow_path)
    # 兼容"日期"列
    if '日期' in df.columns:
        df['时间'] = pd.to_datetime(df['日期'])
    elif '时间' in df.columns:
        df['时间'] = pd.to_datetime(df['时间'])
    else:
        return "资金流向文件缺少日期/时间列。"

    df = df.sort_values('时间')

    # 只保留最近30天
    last_date = df['时间'].max()
    df = df[df['时间'] >= last_date - pd.Timedelta(days=30)]

    # 1. 当日主力净流入/流出金额
    today = df['时间'].dt.date.iloc[-1]
    today_df = df[df['时间'].dt.date == today]
    main_net_today = today_df['主力净流入-净额'].sum() if '主力净流入-净额' in df.columns else None

    # 2. 近5日/10日资金流向趋势
    last_5 = df.groupby(df['时间'].dt.date)['主力净流入-净额'].sum().tail(5) if '主力净流入-净额' in df.columns else None
    last_10 = df.groupby(df['时间'].dt.date)['主力净流入-净额'].sum().tail(10) if '主力净流入-净额' in df.columns else None

    # 3. 资金流向与价格的背离信号
    price_col = '收盘价' if '收盘价' in df.columns else None
    price_trend = None
    fund_trend = None
    divergence = ""
    if price_col and '主力净流入-净额' in df.columns:
        price_trend = df[price_col].iloc[-10:]
        fund_trend = df['主力净流入-净额'].iloc[-10:]
        if price_trend.iloc[-1] > price_trend.iloc[0] and fund_trend.iloc[-1] < fund_trend.iloc[0]:
            divergence = "价格上涨但主力资金流出，警惕背离。"
        elif price_trend.iloc[-1] < price_trend.iloc[0] and fund_trend.iloc[-1] > fund_trend.iloc[0]:
            divergence = "价格下跌但主力资金流入，警惕背离。"
        else:
            divergence = "资金流向与价格趋势一致。"

    # 4. 主力/大单/中单/小单流向分布
    dist_lines = []
    for col in ['超大单净流入-净额', '大单净流入-净额', '中单净流入-净额', '小单净流入-净额']:
        if col in df.columns:
            dist_lines.append(f"{col}\t{df[col].sum():,.0f}")
    if not dist_lines:
        dist_lines = ["无详细分单数据。"]

    # 汇总结论（格式化输出）
    lines = []
    lines.append("资金流向分析")
    if main_net_today is not None:
        lines.append(f"当日主力净流入: {main_net_today:,.0f} 元")
    if last_5 is not None:
        lines.append("近5日主力净流入: ")
        for d, v in last_5.items():
            lines.append(f"{d}: \t{v:,.0f}, ")
    if last_10 is not None:
        lines.append("近10日主力净流入: ")
        for d, v in last_10.items():
            lines.append(f"{d}\t{v:,.0f}")
    lines.append("主力/大单/中单/小单流向分布: ")
    for l in dist_lines:
        lines.append(l)
    lines.append(f"资金流向与价格背离信号: {divergence}")
    return "\n".join(lines)


def analyze_trend(period_dfs: dict, symbol: str, stock_name: str, market: str, date_str: str, output_dir: str):
    doc = Document()
    periods = ['5min', '15min', '30min', '1h', '2h', '1d']
    periods_str = f"分析周期：{', '.join(periods)}（7天）"
    doc.add_paragraph(periods_str)
    doc.add_paragraph("")  # 空行
    doc.add_heading(f"趋势分析报告", 0)
    doc.add_paragraph(f"股票名称: {stock_name}")
    doc.add_paragraph(f"股票代码: {symbol}")
    doc.add_paragraph(f"市场: {market}")
    doc.add_paragraph(f"日期: {date_str}")
    doc.add_paragraph("")

    # ----------- 修正：确保变量总是定义 -----------
    report_lines = []
    fundflow_day = None
    fundflow_hour = None
    fund_flow_path = None

    # 自动查找资金流向文件
    for root, dirs, files in os.walk(output_dir):
        for file in files:
            if symbol in file and "fund_flow" in file and date_str in file:
                fund_flow_path = os.path.join(root, file)
                break

    if fund_flow_path:
        try:
            xls = pd.ExcelFile(fund_flow_path)
            if '1d_fund_flow' in xls.sheet_names:
                fundflow_day = xls.parse('1d_fund_flow')
            if '1h_fund_flow' in xls.sheet_names:
                fundflow_hour = xls.parse('1h_fund_flow')
        except Exception as e:
            fundflow_day = None
            fundflow_hour = None

    # 资金流共振结论
    if fundflow_day is not None and fundflow_hour is not None:
        resonance_result = detect_fund_flow_resonance(fundflow_hour, fundflow_day)
        report_lines.append(f"资金流共振结论：{resonance_result}\n")
        doc.add_paragraph(f"资金流共振结论：{resonance_result}")
    else:
        report_lines.append("资金流共振结论：无数据\n")
        doc.add_paragraph("资金流共振结论：无数据")
    doc.add_paragraph("")  # 新增：空行

    # 分析周期描述
    report_lines.append(periods_str)

    # 最近5个交易日的MA交叉信号和价格位置分析
    doc.add_heading("最近5个交易日趋势分析", level=1)
    for period in ['1d']:  # 只分析日线数据
        df = period_dfs.get(period)
        if df is None or not isinstance(df, pd.DataFrame):
            continue
            
        # 获取最近5个交易日数据
        df = df.tail(5).copy()
        df['时间'] = pd.to_datetime(df['时间'])
        
        # 计算MA交叉信号
        ma_pairs = [(5, 10), (10, 20), (20, 30), (30, 60)]
        for short_period, long_period in ma_pairs:
            short_ma = f'MA{short_period}'
            long_ma = f'MA{long_period}'
            
            if short_ma in df.columns and long_ma in df.columns:
                # 计算金叉死叉
                df[f'{short_ma}_{long_ma}_SIGNAL'] = 0
                df.loc[df[short_ma] > df[long_ma], f'{short_ma}_{long_ma}_SIGNAL'] = 1  # 多头排列
                df.loc[df[short_ma] < df[long_ma], f'{short_ma}_{long_ma}_SIGNAL'] = -1  # 空头排列
                
                # 计算交叉信号
                df[f'{short_ma}_{long_ma}_CROSS'] = 0
                # 上穿信号
                df.loc[(df[short_ma] > df[long_ma]) & (df[short_ma].shift(1) <= df[long_ma].shift(1)), f'{short_ma}_{long_ma}_CROSS'] = 1
                # 下穿信号
                df.loc[(df[short_ma] < df[long_ma]) & (df[short_ma].shift(1) >= df[long_ma].shift(1)), f'{short_ma}_{long_ma}_CROSS'] = -1
        
        # 计算支撑压力位
        window = 20
        threshold = 0.02
        df['SUPPORT'] = df['收盘'].rolling(window=window, center=True).min()
        df['RESISTANCE'] = df['收盘'].rolling(window=window, center=True).max()
        df['PRICE_POSITION'] = 0
        
        # 判断价格位置
        near_support = (df['收盘'] - df['SUPPORT']) / df['SUPPORT'] < threshold
        near_resistance = (df['RESISTANCE'] - df['收盘']) / df['收盘'] < threshold
        df.loc[near_support, 'PRICE_POSITION'] = 1  # 接近支撑位
        df.loc[near_resistance, 'PRICE_POSITION'] = -1  # 接近压力位
        
        # 输出分析结果
        doc.add_paragraph(f"\n{period}周期分析：")
        
        # 输出MA交叉信号
        doc.add_paragraph("MA交叉信号：")
        has_cross_signals = False
        for short_period, long_period in ma_pairs:
            signal_col = f'MA{short_period}_{long_period}_SIGNAL'
            cross_col = f'MA{short_period}_{long_period}_CROSS'
            
            if signal_col in df.columns and cross_col in df.columns:
                for idx, row in df.iterrows():
                    date = row['时间'].strftime('%Y-%m-%d')
                    signal = "多头排列" if row[signal_col] == 1 else "空头排列" if row[signal_col] == -1 else "中性"
                    cross = "金叉" if row[cross_col] == 1 else "死叉" if row[cross_col] == -1 else "无交叉"
                    
                    if row[cross_col] != 0:  # 只输出有交叉信号的日期
                        has_cross_signals = True
                        p = doc.add_paragraph()
                        p.add_run(f"{date}: MA{short_period}/{long_period} {cross}，当前{signal}")
                        if row[cross_col] == 1:  # 金叉用红色
                            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        elif row[cross_col] == -1:  # 死叉用绿色
                            p.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    else:
                        # 输出当前MA排列状态
                        p = doc.add_paragraph()
                        p.add_run(f"{date}: MA{short_period}/{long_period} 当前{signal}")
                        if row[signal_col] == 1:  # 多头排列用红色
                            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        elif row[signal_col] == -1:  # 空头排列用绿色
                            p.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        
        if not has_cross_signals:
            doc.add_paragraph("最近5个交易日无MA交叉信号")
        
        # 输出价格位置
        doc.add_paragraph("\n价格位置：")
        has_position_signals = False
        for idx, row in df.iterrows():
            date = row['时间'].strftime('%Y-%m-%d')
            position = "接近支撑位" if row['PRICE_POSITION'] == 1 else "接近压力位" if row['PRICE_POSITION'] == -1 else "中性区域"
            if row['PRICE_POSITION'] != 0:  # 只输出有明确位置的日期
                has_position_signals = True
                p = doc.add_paragraph()
                p.add_run(f"{date}: {position} (支撑位: {row['SUPPORT']:.2f}, 压力位: {row['RESISTANCE']:.2f})")
                if row['PRICE_POSITION'] == 1:  # 支撑位用红色
                    p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif row['PRICE_POSITION'] == -1:  # 压力位用绿色
                    p.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            else:
                # 输出中性区域的价格位置
                p = doc.add_paragraph()
                p.add_run(f"{date}: {position} (支撑位: {row['SUPPORT']:.2f}, 压力位: {row['RESISTANCE']:.2f})")
        
        if not has_position_signals:
            doc.add_paragraph("最近5个交易日价格处于中性区域")

    # 技术指标分析
    doc.add_heading("技术指标分析", level=1)
    for period in periods:
        df = period_dfs.get(period)
        if df is None or not isinstance(df, pd.DataFrame):
            continue
        doc.add_paragraph(f"周期: {period}")
        macd_signals = filter_recent_signals(detect_macd_divergence(df), df)
        kdj_signals = filter_recent_signals(detect_kdj_divergence(df), df)
        rsi_signals = filter_recent_signals(detect_rsi_divergence(df), df)
        ma_signals = filter_recent_signals(detect_ma_divergence(df), df)
        volume_signals = filter_recent_signals(detect_volume_breakout(df), df)
        add_signals_to_docx(doc, macd_signals, "MACD背离")
        add_signals_to_docx(doc, kdj_signals, "KDJ背离")
        add_signals_to_docx(doc, rsi_signals, "RSI背离")
        add_signals_to_docx(doc, ma_signals, "MA背离")
        add_signals_to_docx(doc, volume_signals, "成交量突破")
        doc.add_paragraph("")

    # 资金流向分析
    doc.add_heading("资金流向分析", level=1)
    if fund_flow_path:
        fund_flow_summary = analyze_fund_flow(fund_flow_path)
        # 新增：分行处理
        for line in fund_flow_summary.splitlines():
            if "资金流向与价格背离信号: 趋势不一致" in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                doc.add_paragraph(line)
        report_lines.append(fund_flow_summary)
    else:
        doc.add_paragraph("未找到资金流向数据文件。")
        report_lines.append("未找到资金流向数据文件。")

    # 生成趋势分析结论
    conclusion = generate_trend_conclusion(period_dfs, fund_flow_path)
    doc.add_heading("趋势分析结论", level=1)
    doc.add_paragraph(conclusion)
    report_lines.append("\n" + conclusion)

    # 组装最终报告文本
    final_report = "\n".join(report_lines)

    file_name = f"{symbol}_{date_str}_trend_report.docx"
    trend_file = os.path.join(output_dir, file_name)
    doc.save(trend_file)
    print(f"✅ 趋势报告生成完成：{trend_file}")

    return final_report


def build_period_dfs(excel_path):
    """
    读取一个多sheet的Excel文件，返回 {周期名: DataFrame} 的字典
    """
    xls = pd.ExcelFile(excel_path)
    period_dfs = {}
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        # 只保留有时间列的周期
        if '时间' in df.columns:
            period_dfs[sheet] = df
    return period_dfs


def batch_trend_analysis(proceeded_dir):
    """
    对 proceeded 目录下所有已处理的Excel文件，批量做趋势分析。
    每个股票只生成一份趋势报告。
    """
    processed_stocks = set()  # 用于记录已处理的股票
    
    for root, dirs, files in os.walk(proceeded_dir):
        for file in files:
            if file.endswith('.xlsx'):
                file_path = os.path.join(root, file)
                m = re.match(r'([A-Za-z0-9\.]+)_([A-Z]+)_.*?(\d{8})', file)
                if m:
                    symbol = m.group(1)
                    market = m.group(2)
                    date_str = m.group(3)
                    
                    # 如果这个股票已经处理过，跳过
                    stock_key = f"{symbol}_{market}_{date_str}"
                    if stock_key in processed_stocks:
                        continue
                    processed_stocks.add(stock_key)
                    
                    stock_name = ""  # 如有股票名映射可查表，否则留空
                    try:
                        print(f"开始分析: {file}")
                        period_dfs = build_period_dfs(file_path)
                        # 修正：按市场分目录
                        trend_dir = os.path.join(proceeded_dir, market)
                        os.makedirs(trend_dir, exist_ok=True)
                        analyze_trend(
                            period_dfs=period_dfs,
                            symbol=symbol,
                            stock_name=stock_name,
                            market=market,
                            date_str=date_str,
                            output_dir=trend_dir
                        )
                    except Exception as e:
                        print(f"趋势分析失败: {file_path}，原因: {e}")


def detect_fund_flow_resonance(fundflow_hour, fundflow_day):
    """
    判断资金流向是否存在共振，返回方向或分歧。
    :param fundflow_hour: 小时级别资金流 dataframe，要求包含 ['主力净流入-净额', '超大单净流入-净额', '大单净流入-净额']
    :param fundflow_day: 日线级别资金流 dataframe，要求同上
    :return: str - ['主力流入共振', '主力流出共振', '多周期资金分歧', '无明显信号']
    """
    try:
        day_main = fundflow_day['主力净流入-净额'].iloc[-1]
        hour_main = fundflow_hour['主力净流入-净额'].tail(3)

        day_big = fundflow_day[['超大单净流入-净额', '大单净流入-净额']].iloc[-1].sum()
        hour_big = fundflow_hour[['超大单净流入-净额', '大单净流入-净额']].tail(3).sum().sum()

        # 共振向上
        if day_main > 0 and hour_main.gt(0).sum() >= 2 and day_big > 0 and hour_big > 0:
            return "主力流入共振"

        # 共振向下
        elif day_main < 0 and hour_main.lt(0).sum() >= 2 and day_big < 0 and hour_big < 0:
            return "主力流出共振"

        # 分歧判断
        elif (day_main > 0 and hour_main.lt(0).sum() >= 2) or (day_main < 0 and hour_main.gt(0).sum() >= 2):
            return "多周期资金分歧"

        else:
            return "无明显信号"
    
    except Exception as e:
        return f"资金流判断失败: {e}"


def detect_volume_breakout(df: pd.DataFrame, 
                         lookback_bars: int = 20,  # 回溯K线数量，用于计算最高价
                         volume_ratio_ma5: float = 1.5,  # 成交量相对5日均量倍数阈值
                         volume_ratio_ma10: float = 2.0,  # 成交量相对10日均量倍数阈值
                         price_increase_ratio: float = 0.01,  # 价格涨幅阈值，默认1%
                         body_ratio_threshold: float = 0.6):  # K线实体占比阈值，默认60%
    """
    检测成交量突破信号，独立判断每个条件并提示
    """
    signals = []
    if len(df) < max(lookback_bars, 10):  # 确保有足够的数据计算MA
        return signals
        
    # 计算成交量移动平均
    df['volume_ma5'] = df['成交量'].rolling(window=5).mean()
    df['volume_ma10'] = df['成交量'].rolling(window=10).mean()
    
    for i in range(lookback_bars, len(df)):
        # 检查是否为无效K线
        if df['最高'].iloc[i] == df['最低'].iloc[i]:
            continue
            
        # 1. 判断是否突破前N日高点
        lookback_high = df['最高'].iloc[i-lookback_bars:i].max()
        price_break = df['收盘'].iloc[i] > lookback_high
        
        # 2. 判断是否放量突破
        volume_break_ma5 = df['成交量'].iloc[i] > df['volume_ma5'].iloc[i] * volume_ratio_ma5
        volume_break_ma10 = df['成交量'].iloc[i] > df['volume_ma10'].iloc[i] * volume_ratio_ma10
        
        # 3. 判断是否为带量长阳K线
        price_increase = df['收盘'].iloc[i] > df['开盘'].iloc[i] * (1 + price_increase_ratio)
        high_low_diff = df['最高'].iloc[i] - df['最低'].iloc[i]
        if high_low_diff > 0:  # 保护除零错误
            body_ratio = abs(df['收盘'].iloc[i] - df['开盘'].iloc[i]) / high_low_diff
            strong_candle = body_ratio >= body_ratio_threshold
        else:
            body_ratio = 0
            strong_candle = False
            
        # 生成信号详情
        details = {
            'price_break': f"突破{lookback_bars}日高点 {lookback_high:.2f}",
            'volume_ratio_ma5': f"{df['成交量'].iloc[i] / df['volume_ma5'].iloc[i]:.1f}倍",
            'volume_ratio_ma10': f"{df['成交量'].iloc[i] / df['volume_ma10'].iloc[i]:.1f}倍",
            'price_increase': f"{(df['收盘'].iloc[i] / df['开盘'].iloc[i] - 1) * 100:.1f}%",
            'body_ratio': f"{body_ratio * 100:.1f}%"
        }
        
        # 独立判断每个条件并生成信号
        if price_break:
            signals.append((
                df['时间'].iloc[i],
                '价格突破',
                "强" if price_increase else "弱",
                f"突破{lookback_bars}日高点 {lookback_high:.2f}，涨幅{details['price_increase']}",
                details
            ))
            
        if volume_break_ma5 or volume_break_ma10:
            volume_type = "放量突破(强)" if volume_break_ma10 else "放量突破(弱)"
            signals.append((
                df['时间'].iloc[i],
                volume_type,
                "强" if volume_break_ma10 else "弱",
                f"成交量放大{details['volume_ratio_ma5']}，10日均量{details['volume_ratio_ma10']}倍",
                details
            ))
            
        if price_increase:
            signals.append((
                df['时间'].iloc[i],
                '长阳K线',
                "强" if strong_candle else "弱",
                f"涨幅{details['price_increase']}，K线实体占比{details['body_ratio']}",
                details
            ))
            
    return signals


def generate_trend_conclusion(period_dfs: dict, fund_flow_path: str = None) -> str:
    # 定义均线周期对
    ma_pairs = [(5, 10), (10, 20), (20, 30), (30, 60)]
    
    # 初始化技术指标信号字典
    tech_signals = {
        '5min': {'macd': [], 'kdj': [], 'volume': []},
        '15min': {'macd': [], 'kdj': [], 'volume': []},
        '30min': {'macd': [], 'kdj': [], 'volume': []},
        '1h': {'macd': [], 'kdj': [], 'volume': []},
        '2h': {'macd': [], 'kdj': [], 'volume': []},
        '1d': {'macd': [], 'kdj': [], 'volume': []}
    }
    
    # 收集各周期技术指标信号
    for period in ['5min', '15min', '30min', '1h', '2h', '1d']:
        df = period_dfs.get(period)
        if df is not None:
            tech_signals[period]['macd'] = filter_recent_signals(detect_macd_divergence(df), df)
            tech_signals[period]['kdj'] = filter_recent_signals(detect_kdj_divergence(df), df)
            tech_signals[period]['volume'] = filter_recent_signals(detect_volume_breakout(df), df)
    
    conclusion = []
    
    # 1. 趋势分析结论（提前到最前面）
    conclusion.append("趋势分析结论：")
    
    # 1.1 主要特征（作为结论的核心部分）
    features = []
    
    # 短期背离信号
    short_term_divergence = sum(len(tech_signals[period]['macd']) + len(tech_signals[period]['kdj']) 
                              for period in ['5min', '15min', '30min', '1h'])
    if short_term_divergence > 0:
        features.append(f"短期（5分钟到1小时）出现{short_term_divergence}次技术指标背离信号")
    
    # 成交量突破信号
    volume_signals_count = sum(len(tech_signals[period]['volume']) 
                             for period in ['5min', '15min', '30min', '1h'])
    if volume_signals_count > 0:
        features.append(f"成交量突破信号主要集中在短期周期，共{volume_signals_count}次")
    
    # 资金流向特征
    if fund_flow_path:
        fund_flow_summary = analyze_fund_flow(fund_flow_path)
        if "主力净流入: -" in fund_flow_summary:
            features.append("资金流向持续为负，显示主力资金持续流出")
        if "小单净流入" in fund_flow_summary and "主力净流入: -" in fund_flow_summary:
            features.append("小单资金大幅流入，与主力资金形成明显对比")
    
    # 输出主要特征作为结论
    if features:
        conclusion.extend([f"- {feature}" for feature in features])
    else:
        conclusion.append("- 当前无明显趋势特征")
    
    # 1.2 风险提示（作为结论的补充部分）
    risk_warnings = []
    
    # 检查短期背离信号
    if short_term_divergence >= 5:
        risk_warnings.append("短期技术指标出现多次顶背离，需警惕回调风险")
    
    # 检查资金流向
    if fund_flow_path and "主力净流入: -" in fund_flow_summary:
        risk_warnings.append("主力资金持续流出，累计金额较大")
    
    # 检查价格位置
    df_1d = period_dfs.get('1d')
    if df_1d is not None and 'PRICE_POSITION' in df_1d.columns:
        if all(df_1d['PRICE_POSITION'] == 0):
            risk_warnings.append("价格处于中性区域，缺乏明确方向性")
    
    if risk_warnings:
        conclusion.append("\n风险提示：")
        conclusion.extend([f"- {warning}" for warning in risk_warnings])
    
    # 2. 分析依据
    conclusion.append("\n分析依据：")
    
    # 2.1 日线级别分析
    conclusion.append("\n日线级别分析：")
    df_1d = period_dfs.get('1d')
    if df_1d is not None:
        # MA交叉信号
        has_cross_signals = False
        for ma1, ma2 in ma_pairs:
            if f'MA{ma1}_CROSS_MA{ma2}' in df_1d.columns:
                cross_signals = df_1d[f'MA{ma1}_CROSS_MA{ma2}'].iloc[-5:].sum()
                if cross_signals > 0:
                    has_cross_signals = True
                    break
        
        conclusion.append("MA交叉信号：" + ("最近5个交易日无MA交叉信号" if not has_cross_signals else "..."))
        
        # 价格位置
        if 'PRICE_POSITION' in df_1d.columns:
            if all(df_1d['PRICE_POSITION'] == 0):
                conclusion.append("价格位置：所有交易日都处于中性区域")
            else:
                support_days = sum(df_1d['PRICE_POSITION'] == 1)
                resistance_days = sum(df_1d['PRICE_POSITION'] == -1)
                if support_days > 0:
                    conclusion.append(f"价格位置：{support_days}个交易日接近支撑位")
                if resistance_days > 0:
                    conclusion.append(f"价格位置：{resistance_days}个交易日接近压力位")
        else:
            conclusion.append("价格位置：数据不可用")
            
        # 支撑位和压力位
        if 'SUPPORT' in df_1d.columns and 'RESISTANCE' in df_1d.columns:
            if not pd.isna(df_1d['SUPPORT'].iloc[-1]) and not pd.isna(df_1d['RESISTANCE'].iloc[-1]):
                conclusion.append(f"支撑位和压力位：支撑位 {df_1d['SUPPORT'].iloc[-1]:.2f}，压力位 {df_1d['RESISTANCE'].iloc[-1]:.2f}")
            else:
                conclusion.append("支撑位和压力位数据缺失（显示为nan）")
        else:
            conclusion.append("支撑位和压力位：数据不可用")
    
    # 2.2 技术指标分析（按周期）
    conclusion.append("\n技术指标分析（按周期）：")
    for period in ['5min', '15min', '30min', '1h', '2h', '1d']:
        period_signals = []
        df = period_dfs.get(period)
        if df is not None:
            # MACD背离
            if tech_signals[period]['macd']:
                top_div = sum(1 for s in tech_signals[period]['macd'] if '顶背离' in s[1])
                bottom_div = sum(1 for s in tech_signals[period]['macd'] if '底背离' in s[1])
                if top_div > 0 and bottom_div > 0:
                    period_signals.append(f"MACD背离：顶背离{top_div}次，底背离{bottom_div}次")
                elif top_div > 0:
                    period_signals.append(f"MACD顶背离：{top_div}次")
                elif bottom_div > 0:
                    period_signals.append(f"MACD底背离：{bottom_div}次")
            
            # KDJ背离
            if tech_signals[period]['kdj']:
                top_div = sum(1 for s in tech_signals[period]['kdj'] if '顶背离' in s[1])
                bottom_div = sum(1 for s in tech_signals[period]['kdj'] if '底背离' in s[1])
                if top_div > 0 and bottom_div > 0:
                    period_signals.append(f"KDJ背离：顶背离{top_div}次，底背离{bottom_div}次")
                elif top_div > 0:
                    period_signals.append(f"KDJ顶背离：{top_div}次")
                elif bottom_div > 0:
                    period_signals.append(f"KDJ底背离：{bottom_div}次")
            
            # 成交量突破
            if tech_signals[period]['volume']:
                strong_volume = sum(1 for s in tech_signals[period]['volume'] if s[2] == "强")
                weak_volume = sum(1 for s in tech_signals[period]['volume'] if s[2] == "弱")
                if strong_volume > 0 and weak_volume > 0:
                    period_signals.append(f"成交量突破：出现{strong_volume + weak_volume}次，其中{strong_volume}次为强信号")
                elif strong_volume > 0:
                    period_signals.append(f"成交量突破：出现{strong_volume}次，均为强信号")
                elif weak_volume > 0:
                    period_signals.append(f"成交量突破：出现{weak_volume}次，均为弱信号")
            
            if period_signals:
                conclusion.append(f"{chr(97 + ['5min', '15min', '30min', '1h', '2h', '1d'].index(period))}) {period}周期：")
                conclusion.extend([f"   {signal}" for signal in period_signals])
            else:
                conclusion.append(f"{chr(97 + ['5min', '15min', '30min', '1h', '2h', '1d'].index(period))}) {period}周期：")
                conclusion.append("   所有技术指标均无信号")
    
    # 2.3 资金流向分析
    conclusion.append("\n资金流向分析：")
    if fund_flow_path:
        fund_flow_summary = analyze_fund_flow(fund_flow_path)
        fund_lines = fund_flow_summary.split('\n')
        for line in fund_lines:
            if any(x in line for x in ['当日主力净流入', '近5日主力净流入', '近10日主力净流入', '超大单', '大单', '中单', '小单']):
                conclusion.append(line)
    
    return "\n".join(conclusion)
